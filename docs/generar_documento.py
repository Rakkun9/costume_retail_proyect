from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Márgenes APA (2.54 cm = 1 in) ──────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

# ── Helpers ─────────────────────────────────────────────────────────────────
def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold  = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para(text="", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=False,
         italic=False, spacing_before=0, spacing_after=0, first_indent=0,
         line_spacing=None, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(spacing_before)
    pf.space_after  = Pt(spacing_after)
    if first_indent:
        pf.first_line_indent = Cm(first_indent)
    if line_spacing:
        from docx.shared import Pt as Ptt
        pf.line_spacing = Ptt(line_spacing)
    else:
        from docx.enum.text import WD_LINE_SPACING
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def heading(text, level=1, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    from docx.enum.text import WD_LINE_SPACING
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    return p

def add_page_break():
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ════════════════════════════════════════════════════════════════════════════
for _ in range(4):
    para(line_spacing=12)

para("Ucompensar", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
para("Facultad de Ingeniería", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
para("Asignatura: Patrones de Diseño", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

for _ in range(4):
    para(line_spacing=12)

para(
    "Aplicación de Patrones de Diseño en el Sistema de\n"
    "Alquiler de Prendas de Disfraz (Costume Retail)",
    align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True
)

for _ in range(4):
    para(line_spacing=12)

para("Integrantes:", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
para("Juan Nicolay Buitrago Vargas", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
para("Santiago Bayona Bohórquez",   align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

for _ in range(3):
    para(line_spacing=12)

para("Docente:", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
para("Oswaldo Pérez Murillo",        align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

for _ in range(3):
    para(line_spacing=12)

para("Bogotá D.C., Colombia", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
para("Marzo de 2026",          align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  RESUMEN (Abstract en APA)
# ════════════════════════════════════════════════════════════════════════════
heading("Resumen", level=1, size=12)
para(
    "El presente documento describe la identificación y aplicación de patrones de diseño "
    "de software en el desarrollo del sistema Costume Retail, una aplicación web para la "
    "gestión del servicio de alquiler de prendas de disfraz. El sistema fue construido "
    "utilizando NestJS como framework backend y Astro como frontend, con una base de datos "
    "PostgreSQL alojada en Supabase. A lo largo del desarrollo se identificaron y aplicaron "
    "seis patrones de diseño fundamentales: Singleton, Repository, Facade, Decorator, "
    "Factory Method y Module. Para cada patrón se presenta su nombre, tipo, justificación "
    "de uso y descripción detallada de su implementación en el contexto del proyecto. "
    "La correcta aplicación de estos patrones contribuyó a una arquitectura modular, "
    "mantenible y escalable.",
    first_indent=1.27, size=12
)
p_kw = doc.add_paragraph()
p_kw.alignment = WD_ALIGN_PARAGRAPH.LEFT
from docx.enum.text import WD_LINE_SPACING
p_kw.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
r1 = p_kw.add_run("Palabras clave: ")
set_font(r1, bold=True, italic=True, size=12)
r2 = p_kw.add_run("patrones de diseño, NestJS, Singleton, Repository, Facade, Decorator, Factory Method, arquitectura de software.")
set_font(r2, italic=True, size=12)

add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  INTRODUCCIÓN
# ════════════════════════════════════════════════════════════════════════════
heading("Introducción", level=1, size=12)
intro_pars = [
    ("    Los patrones de diseño representan soluciones reutilizables a problemas recurrentes "
     "en el desarrollo de software. Fueron formalizados por Gamma et al. (1994) en su obra "
     "seminal Design Patterns: Elements of Reusable Object-Oriented Software, conocida "
     "popularmente como el 'Gang of Four' (GoF). Desde entonces, su adopción se ha "
     "extendido a todos los paradigmas y tecnologías del desarrollo moderno."),
    ("    El proyecto Costume Retail es un sistema web para la gestión de alquiler de prendas "
     "de disfraz, desarrollado como trabajo académico en la asignatura Patrones de Diseño de "
     "la Universidad Compensar. El sistema contempla el registro de prendas, clientes y "
     "empleados; la creación de servicios de alquiler con validaciones de disponibilidad; "
     "consultas por número, cliente y fecha; y la gestión de una cola de lavandería con "
     "prioridad. El backend fue implementado con NestJS (Node.js), TypeORM y PostgreSQL "
     "a través de Supabase, mientras que el frontend se desarrolló con Astro y Bootstrap 5."),
    ("    En el proceso de diseño y desarrollo del sistema se identificaron y aplicaron "
     "seis patrones de diseño, seleccionados según las necesidades arquitecturales del "
     "proyecto: Singleton, Repository, Facade, Decorator, Factory Method y Module. "
     "El presente documento expone en detalle cada uno de estos patrones, su clasificación, "
     "la justificación de su uso y la forma concreta en que se implementaron en el código "
     "fuente del sistema.")
]
for text in intro_pars:
    para(text, size=12)

add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  TABLA RESUMEN DE PATRONES
# ════════════════════════════════════════════════════════════════════════════
heading("Tabla de Patrones Identificados", level=1, size=12)
para(
    "    A continuación se presenta la tabla resumen de los patrones de diseño "
    "identificados en el sistema Costume Retail, con su clasificación y justificación.",
    size=12
)
para(size=12)

# Nota de tabla
p_note = doc.add_paragraph()
p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
from docx.enum.text import WD_LINE_SPACING
p_note.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
r = p_note.add_run("Tabla 1")
set_font(r, bold=True, italic=True, size=12)

p_note2 = doc.add_paragraph()
p_note2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_note2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
r2 = p_note2.add_run("Patrones de diseño aplicados en el sistema Costume Retail")
set_font(r2, italic=True, size=12)

# Tabla
headers = ["N°", "Nombre del Patrón", "Tipo", "Justificación de uso"]
rows = [
    ["1", "Singleton",      "Creacional",
     "NestJS registra cada proveedor (service) como instancia única por módulo, "
     "garantizando un único punto de acceso a los repositorios de base de datos."],
    ["2", "Repository",     "Estructural",
     "TypeORM expone el patrón Repository para aislar el acceso a la base de datos, "
     "manteniendo la lógica de negocio independiente de la capa de persistencia."],
    ["3", "Facade",         "Estructural",
     "Cada Service actúa como fachada que simplifica las operaciones complejas "
     "(validaciones, consultas relacionadas) ocultando los detalles a los controladores."],
    ["4", "Decorator",      "Estructural",
     "Los decoradores de TypeORM (@Entity, @Column, @PrimaryColumn) y de NestJS "
     "(@Controller, @Injectable, @Get, @Post) añaden metadatos sin modificar las clases."],
    ["5", "Factory Method", "Creacional",
     "El método repository.create(dto) actúa como fábrica que instancia entidades "
     "desde DTOs, desacoplando la creación de objetos de su persistencia."],
    ["6", "Module",         "Estructural",
     "NestJS organiza el sistema en módulos cohesivos (PrendasModule, ClientesModule, "
     "etc.) que encapsulan sus propios controladores, servicios y repositorios."],
]

table = doc.add_table(rows=1 + len(rows), cols=4)
table.style = "Table Grid"

# Header row
widths = [Cm(1), Cm(4), Cm(3.5), Cm(8.5)]
hdr = table.rows[0]
hdr_texts = headers
for i, (cell, txt, w) in enumerate(zip(hdr.cells, hdr_texts, widths)):
    cell.width = w
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p2 = cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run(txt)
    set_font(r, bold=True, size=11)
    # Gray shading
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "D9D9D9")
    tcPr.append(shd)

# Data rows
for row_data in rows:
    row = table.add_row()
    for i, (cell, txt, w) in enumerate(zip(row.cells, row_data, widths)):
        cell.width = w
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p2 = cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 3 else WD_ALIGN_PARAGRAPH.LEFT
        r = p2.add_run(txt)
        set_font(r, size=10)

# Nota al pie de tabla
para(size=12)
p_nt = doc.add_paragraph()
p_nt.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
r_nt = p_nt.add_run("Nota. ")
set_font(r_nt, italic=True, size=10)
r_nt2 = p_nt.add_run(
    "La clasificación Creacional/Estructural sigue la taxonomía de Gamma et al. (1994). "
    "El patrón Module, aunque propio de NestJS, responde al principio de diseño modular "
    "del patrón Facade extendido."
)
set_font(r_nt2, size=10)

add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  DESCRIPCIÓN DETALLADA DE PATRONES
# ════════════════════════════════════════════════════════════════════════════
heading("Descripción Detallada de los Patrones de Diseño", level=1, size=12)

# ── Patrón 1: Singleton ─────────────────────────────────────────────────────
heading("Patrón Singleton", level=2, size=12)

para("Tipo: Creacional", size=12, bold=True)
para(
    "    El patrón Singleton garantiza que una clase tenga una única instancia durante "
    "todo el ciclo de vida de la aplicación y proporciona un punto de acceso global a "
    "dicha instancia (Gamma et al., 1994). Este patrón es especialmente útil cuando se "
    "necesita controlar el acceso a recursos compartidos, como conexiones a bases de datos.",
    size=12
)
para(
    "    En NestJS, el sistema de inyección de dependencias registra cada proveedor con "
    "un alcance (scope) Singleton por defecto. Esto significa que la primera vez que un "
    "módulo solicita un servicio, el framework instancia la clase y la reutiliza en "
    "todas las solicitudes posteriores dentro del mismo contexto de módulo. En el "
    "proyecto Costume Retail, servicios como PrendasService, ClientesService, "
    "EmpleadosService, ServiciosService y LavanderiaService se comportan como "
    "Singletons, garantizando que el repositorio TypeORM que cada uno inyecta sea "
    "la misma instancia compartida.",
    size=12
)

p_code = doc.add_paragraph()
p_code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p_code.paragraph_format.left_indent = Cm(1.27)
r_c = p_code.add_run(
    "@Injectable()  // Scope.DEFAULT = Singleton\n"
    "export class PrendasService {\n"
    "  constructor(\n"
    "    @InjectRepository(Prenda)\n"
    "    private readonly prendasRepository: Repository<Prenda>,\n"
    "  ) {}"
)
set_font(r_c, name="Courier New", size=10)

para(
    "    La anotación @Injectable() registra PrendasService en el contenedor de IoC de "
    "NestJS. Dado que no se especifica un scope diferente, el framework crea una única "
    "instancia del servicio y la reutiliza, evitando el costo de instanciación en cada "
    "petición HTTP y asegurando consistencia en el acceso a los datos.",
    size=12
)

# ── Patrón 2: Repository ────────────────────────────────────────────────────
heading("Patrón Repository", level=2, size=12)

para("Tipo: Estructural", size=12, bold=True)
para(
    "    El patrón Repository actúa como intermediario entre la lógica de negocio y la "
    "capa de acceso a datos, encapsulando las operaciones de consulta y persistencia "
    "detrás de una interfaz orientada al dominio (Fowler, 2002). Su objetivo principal "
    "es desacoplar la lógica de negocio de los detalles de almacenamiento.",
    size=12
)
para(
    "    En Costume Retail, TypeORM implementa este patrón a través de la clase genérica "
    "Repository<T>. Cada servicio del sistema recibe por inyección de dependencias un "
    "repositorio específico para su entidad. Por ejemplo, PrendasService recibe un "
    "Repository<Prenda>, y a través de él realiza todas las operaciones con la tabla "
    "prendas sin escribir SQL directamente.",
    size=12
)
p_code2 = doc.add_paragraph()
p_code2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p_code2.paragraph_format.left_indent = Cm(1.27)
r_c2 = p_code2.add_run(
    "// Uso del repositorio en PrendasService\n"
    "async findByReferencia(referencia: string): Promise<Prenda> {\n"
    "  const prenda = await this.prendasRepository.findOne({\n"
    "    where: { referencia },\n"
    "  });\n"
    "  if (!prenda) throw new NotFoundException(...);\n"
    "  return prenda;\n"
    "}"
)
set_font(r_c2, name="Courier New", size=10)

para(
    "    Esta implementación permite que el servicio opere sobre objetos del dominio "
    "(Prenda, Cliente, etc.) sin conocer los detalles de la conexión a PostgreSQL. "
    "Si en el futuro se cambia el motor de base de datos, únicamente se modifica la "
    "configuración de TypeORM, sin afectar la lógica de negocio.",
    size=12
)

# ── Patrón 3: Facade ────────────────────────────────────────────────────────
heading("Patrón Facade", level=2, size=12)

para("Tipo: Estructural", size=12, bold=True)
para(
    "    El patrón Facade proporciona una interfaz simplificada a un conjunto complejo "
    "de clases, subsistemas o servicios (Gamma et al., 1994). Su propósito es reducir "
    "la complejidad visible para el cliente, delegando el trabajo a los subsistemas "
    "internos según sea necesario.",
    size=12
)
para(
    "    En el proyecto, la capa de servicios funciona como una fachada para los "
    "controladores. El caso más representativo es ServiciosService, cuyo método create() "
    "orquesta internamente: la validación del cliente (ClientesService), la validación "
    "del empleado (EmpleadosService), la verificación de disponibilidad de cada prenda "
    "(PrendasService) y la verificación de conflictos de fecha mediante QueryBuilder de "
    "TypeORM. El controlador ServiciosController simplemente recibe el DTO y delega "
    "toda esa complejidad al servicio con una sola llamada.",
    size=12
)
p_code3 = doc.add_paragraph()
p_code3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p_code3.paragraph_format.left_indent = Cm(1.27)
r_c3 = p_code3.add_run(
    "// ServiciosController actúa solo como punto de entrada\n"
    "@Post()\n"
    "create(@Body() dto: CreateServicioDto) {\n"
    "  return this.serviciosService.create(dto); // Fachada\n"
    "}"
)
set_font(r_c3, name="Courier New", size=10)

para(
    "    El cliente HTTP (o el frontend) no necesita conocer que para registrar un "
    "servicio se consultan tres entidades diferentes y se ejecutan múltiples validaciones. "
    "La fachada (el servicio) abstrae toda esa complejidad.",
    size=12
)

# ── Patrón 4: Decorator ─────────────────────────────────────────────────────
heading("Patrón Decorator", level=2, size=12)

para("Tipo: Estructural", size=12, bold=True)
para(
    "    El patrón Decorator permite añadir comportamientos o responsabilidades a un "
    "objeto de forma dinámica sin modificar su clase base (Gamma et al., 1994). En "
    "TypeScript y JavaScript moderno, los decoradores de clase son una implementación "
    "nativa de este patrón a través de metaprogramación.",
    size=12
)
para(
    "    En Costume Retail, el patrón Decorator está presente en dos niveles. A nivel "
    "de entidades, TypeORM utiliza decoradores para adjuntar metadatos de mapeo "
    "objeto-relacional: @Entity() indica que la clase es una tabla, @Column() y "
    "@PrimaryColumn() definen las columnas, mientras que @ManyToMany(), @OneToMany() "
    "y @ManyToOne() establecen las relaciones. A nivel de la API, NestJS emplea "
    "@Controller(), @Get(), @Post(), @Param(), @Body() y @ApiOperation() de Swagger "
    "para definir rutas, validaciones y documentación sin modificar la lógica de los "
    "métodos.",
    size=12
)
p_code4 = doc.add_paragraph()
p_code4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p_code4.paragraph_format.left_indent = Cm(1.27)
r_c4 = p_code4.add_run(
    "@Entity('prendas')          // Mapea la clase a la tabla 'prendas'\n"
    "export class Prenda {\n"
    "  @PrimaryColumn()          // Columna clave primaria\n"
    "  referencia: string;\n\n"
    "  @Column({ type: 'enum', enum: EstadoPrenda })\n"
    "  estado: EstadoPrenda;     // Columna con tipo enumerado\n\n"
    "  @ManyToMany(() => ServicioAlquiler)\n"
    "  servicios: ServicioAlquiler[];  // Relación N:M\n"
    "}"
)
set_font(r_c4, name="Courier New", size=10)

para(
    "    Gracias a este patrón, la clase Prenda no necesita heredar de ninguna clase "
    "base ni implementar interfaces específicas de TypeORM; los metadatos se adjuntan "
    "de forma no intrusiva, manteniendo las clases limpias y cohesivas.",
    size=12
)

# ── Patrón 5: Factory Method ────────────────────────────────────────────────
heading("Patrón Factory Method", level=2, size=12)

para("Tipo: Creacional", size=12, bold=True)
para(
    "    El patrón Factory Method define una interfaz para crear un objeto, pero delega "
    "a las subclases la decisión sobre qué clase concreta instanciar (Gamma et al., 1994). "
    "Su principal beneficio es desacoplar el proceso de creación de objetos de su uso.",
    size=12
)
para(
    "    En Costume Retail, TypeORM expone el método repository.create(plainObject) que "
    "actúa como fábrica: transforma un objeto plano (generalmente un DTO) en una "
    "instancia de la entidad correspondiente con todos sus metadatos de TypeORM "
    "correctamente configurados, pero sin persistirlo aún en la base de datos. La "
    "persistencia ocurre únicamente al llamar repository.save(entity). Esta separación "
    "entre creación y persistencia es la esencia del Factory Method.",
    size=12
)
p_code5 = doc.add_paragraph()
p_code5.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p_code5.paragraph_format.left_indent = Cm(1.27)
r_c5 = p_code5.add_run(
    "// Factory Method en PrendasService\n"
    "async create(dto: CreatePrendaDto): Promise<Prenda> {\n"
    "  // create() = Factory: instancia Prenda desde el DTO\n"
    "  const prenda = this.prendasRepository.create(dto);\n"
    "  // save() = persiste la instancia ya creada\n"
    "  return this.prendasRepository.save(prenda);\n"
    "}"
)
set_font(r_c5, name="Courier New", size=10)

para(
    "    Este patrón también aparece en ServiciosService, donde se construye el objeto "
    "ServicioAlquiler con sus relaciones (cliente, empleado, prendas) antes de "
    "persistirlo, permitiendo aplicar toda la lógica de validación sobre el objeto "
    "creado en memoria antes de comprometerse con la base de datos.",
    size=12
)

# ── Patrón 6: Module ────────────────────────────────────────────────────────
heading("Patrón Module", level=2, size=12)

para("Tipo: Estructural", size=12, bold=True)
para(
    "    El patrón Module organiza el código en unidades autónomas y cohesivas que "
    "encapsulan un conjunto relacionado de funcionalidades, exponiendo solo lo necesario "
    "hacia el exterior (Osmani, 2012). Es una especialización del principio de encapsulación "
    "y se relaciona con el patrón Facade a nivel arquitectural.",
    size=12
)
para(
    "    NestJS hace del patrón Module un ciudadano de primera clase. En Costume Retail "
    "se crearon cinco módulos de dominio: PrendasModule, ClientesModule, EmpleadosModule, "
    "ServiciosModule y LavanderiaModule. Cada módulo declara sus propios controladores, "
    "servicios y el repositorio TypeORM que necesita. La dependencia entre módulos se "
    "expresa explícitamente mediante imports y exports, de modo que ServiciosModule puede "
    "usar PrendasService únicamente porque PrendasModule lo exporta.",
    size=12
)
p_code6 = doc.add_paragraph()
p_code6.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p_code6.paragraph_format.left_indent = Cm(1.27)
r_c6 = p_code6.add_run(
    "@Module({\n"
    "  imports:     [TypeOrmModule.forFeature([Prenda])],\n"
    "  providers:   [PrendasService],\n"
    "  controllers: [PrendasController],\n"
    "  exports:     [PrendasService],  // Expone el servicio\n"
    "})\n"
    "export class PrendasModule {}"
)
set_font(r_c6, name="Courier New", size=10)

para(
    "    Esta organización modular facilita la escalabilidad del sistema: agregar una "
    "nueva funcionalidad implica crear un nuevo módulo sin afectar los existentes. "
    "Adicionalmente, simplifica las pruebas unitarias al permitir importar únicamente "
    "el módulo bajo prueba en los tests.",
    size=12
)

add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  CONCLUSIONES
# ════════════════════════════════════════════════════════════════════════════
heading("Conclusiones", level=1, size=12)

conclusiones = [
    ("    La aplicación del patrón Singleton en el sistema de inyección de dependencias "
     "de NestJS garantizó un uso eficiente de recursos, al evitar la creación redundante "
     "de instancias de servicios y repositorios en cada solicitud HTTP, mejorando el "
     "rendimiento general de la aplicación."),
    ("    El patrón Repository, implementado a través de TypeORM, logró un desacoplamiento "
     "efectivo entre la lógica de negocio y la capa de persistencia. Esta separación "
     "permite cambiar el motor de base de datos sin impactar los servicios de dominio, "
     "lo cual es fundamental para la mantenibilidad a largo plazo del sistema."),
    ("    El uso del patrón Facade en la capa de servicios redujo significativamente la "
     "complejidad visible para los controladores. El caso más notorio fue ServiciosService, "
     "que orquesta validaciones en múltiples entidades bajo una única interfaz de método, "
     "respetando el principio de responsabilidad única."),
    ("    Los decoradores de TypeORM y NestJS, como implementación del patrón Decorator, "
     "demostraron ser una herramienta poderosa para añadir comportamientos transversales "
     "(mapeo ORM, rutas HTTP, validación, documentación) sin contaminar la lógica de "
     "negocio de las entidades y controladores."),
    ("    La combinación del patrón Factory Method con el Repository permitió separar "
     "limpiamente la creación de entidades de su persistencia, facilitando la validación "
     "de objetos en memoria antes de comprometer cambios en la base de datos."),
    ("    Finalmente, la arquitectura modular de NestJS, alineada con el patrón Module, "
     "demostró ser el marco ideal para aplicar el conjunto de patrones de forma coherente. "
     "La modularidad permitió que cada componente del sistema fuera desarrollado, probado "
     "y razonado de forma independiente, reduciendo el acoplamiento y favoreciendo la "
     "escalabilidad del proyecto Costume Retail."),
]

for c in conclusiones:
    para(c, size=12)

add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  REFERENCIAS (APA 7ma edición)
# ════════════════════════════════════════════════════════════════════════════
heading("Referencias", level=1, size=12)

refs = [
    ("Fowler, M. (2002). ", "Patterns of enterprise application architecture. "
     "Addison-Wesley Professional."),
    ("Gamma, E., Helm, R., Johnson, R., y Vlissides, J. (1994). ", "Design patterns: "
     "Elements of reusable object-oriented software. Addison-Wesley Professional."),
    ("NestJS. (2024). ", "NestJS documentation — providers. "
     "https://docs.nestjs.com/providers"),
    ("NestJS. (2024). ", "NestJS documentation — modules. "
     "https://docs.nestjs.com/modules"),
    ("Osmani, A. (2012). ", "Learning JavaScript design patterns. O'Reilly Media. "
     "https://www.patterns.dev/posts/classic-design-patterns/"),
    ("TypeORM. (2024). ", "TypeORM documentation — repository API. "
     "https://typeorm.io/repository-api"),
    ("Freeman, E., Robson, E., Bates, B., y Sierra, K. (2004). ", "Head first design "
     "patterns. O'Reilly Media."),
]

for author, rest in refs:
    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p_ref.paragraph_format.left_indent = Cm(1.27)
    p_ref.paragraph_format.first_line_indent = Cm(-1.27)
    r_a = p_ref.add_run(author)
    set_font(r_a, size=12)
    r_b = p_ref.add_run(rest)
    set_font(r_b, italic=True, size=12)
    # Ensure URLs not italic
    if "http" in rest:
        # split at url
        idx = rest.find("http")
        body = rest[:idx]
        url  = rest[idx:]
        # remove r_b and re-add
        p_ref.runs[-1].text = body
        r_url = p_ref.add_run(url)
        set_font(r_url, italic=False, size=12, color=(5, 99, 193))

# ── Guardar ──────────────────────────────────────────────────────────────────
out = r"c:\Users\jnico\Desktop\University Projects\costume_retail_proyect\docs\Patrones_Diseño_CostumeRetail.docx"
doc.save(out)
print(f"Documento guardado en: {out}")
